# -*- coding: utf-8 -*-
from flask import Flask
from pdfminer.pdfinterp import PDFResourceManager, PDFPageInterpreter
from pdfminer.converter import TextConverter
from pdfminer.layout import LAParams
from pdfminer.pdfpage import PDFPage
from io import BytesIO
import re
import json
from docx import Document
import os
import sys
reload(sys)
sys.setdefaultencoding("utf-8")

app = Flask(__name__)

@app.route('/')
def convert_pdf(path='provide path here', format='text', codec='utf-8'):
    rsrcmgr = PDFResourceManager()
    retstr = BytesIO()
    laparams = LAParams()
    if format == 'text':
        device = TextConverter(rsrcmgr, retstr, codec=codec, laparams=laparams)
    else:
        raise ValueError('Please provide the format to extract')
    fp = open(path, 'rb')
    interpreter = PDFPageInterpreter(rsrcmgr, device)
    maxpages = 500 #mention the maximum pages here (Note: Large number of pages will decrease the performance.)
    caching = True
    page_numbers=set()
    for page in PDFPage.get_pages(fp, page_numbers, maxpages=maxpages,caching=caching, check_extractable=True):
        interpreter.process_page(page)
    text = retstr.getvalue().decode()
    fp.close()
    device.close()
    retstr.close()
    bulletins_data = re.findall('•([^•]+)*', str(text))
    list_of_bullet_points = []
    json_dict = {}
    for points in bulletins_data:
        list_of_bullet_points.append(points)
    json_dict['bulletins'] = list_of_bullet_points
    json_data= json.dumps(json_dict)
    parsed = json.loads(json_data)
    final_data = json.dumps(parsed, indent=4, sort_keys=True) #creates a pretty json with the data extracted
    document = Document()  # creates a new document
    document.add_heading('Bulletins data in the PDF')
    document.add_paragraph(str(final_data))
    document.save('json_data.docx')  # saves it to the filesystem
    os.startfile("json_data.docx")  # will open the file
    return ''

if __name__ == '__main__':
    app.run(debug=True)
